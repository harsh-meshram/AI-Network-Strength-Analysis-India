"""
Convert Markdown reports to .docx files using python-docx.
Parses Markdown structure (headings, tables, bullet points, code blocks) into Word format.
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def parse_md_to_docx(md_path, docx_path, title):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i].rstrip('\n').rstrip('\r')
        
        # Skip empty metadata lines at top (bold/italic markers)
        if line.startswith('---'):
            i += 1
            continue
        
        # Code block handling
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block - write accumulated code
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(30, 30, 30)
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                code_lines = []
                in_code_block = False
            else:
                # Flush any pending table
                if in_table and table_rows:
                    write_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Table handling
        if '|' in line and line.strip().startswith('|'):
            # Check if separator line
            stripped = line.strip()
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                i += 1
                continue
            # Parse table row
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if not in_table:
                in_table = True
            table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table and table_rows:
                write_table(doc, table_rows)
                table_rows = []
                in_table = False
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Headings
        if line.startswith('# '):
            text = clean_md(line[2:].strip())
            p = doc.add_heading(text, level=0)
            i += 1
            continue
        elif line.startswith('## '):
            text = clean_md(line[3:].strip())
            doc.add_heading(text, level=1)
            i += 1
            continue
        elif line.startswith('### '):
            text = clean_md(line[4:].strip())
            doc.add_heading(text, level=2)
            i += 1
            continue
        elif line.startswith('#### '):
            text = clean_md(line[5:].strip())
            doc.add_heading(text, level=3)
            i += 1
            continue
        
        # Blockquote
        if line.strip().startswith('> '):
            text = clean_md(line.strip()[2:])
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].italic = True if p.runs else None
            i += 1
            continue
        
        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = clean_md(line.strip()[2:])
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue
        
        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.*)', line.strip())
        if m:
            text = clean_md(m.group(2))
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            i += 1
            continue
        
        # Regular paragraph
        text = clean_md(line.strip())
        if text:
            p = doc.add_paragraph()
            add_formatted_text(p, text)
        
        i += 1
    
    # Flush remaining table
    if in_table and table_rows:
        write_table(doc, table_rows)
    
    doc.save(docx_path)
    print(f"Created: {docx_path}")


def write_table(doc, rows):
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = clean_md(cell_text)
                # Bold header row
                if r_idx == 0:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(10)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)


def clean_md(text):
    """Remove markdown formatting characters."""
    # Remove bold
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Remove italic  
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # Remove inline code
    text = re.sub(r'`(.*?)`', r'\1', text)
    # Remove links [text](url)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    # Remove images
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    return text.strip()


def add_formatted_text(paragraph, text):
    """Add text to paragraph with basic bold/italic support."""
    # Simple approach: just add as plain text (formatting removed by clean_md)
    paragraph.add_run(text)


# Convert both reports
base = r'H:\AI Network Strength Analysis India'

parse_md_to_docx(
    os.path.join(base, 'AntiGravity_NonTechnical_Status.md'),
    os.path.join(base, 'AntiGravity_NonTechnical_Status.docx'),
    'AntiGravity — Non-Technical Project Status Report'
)

parse_md_to_docx(
    os.path.join(base, 'AntiGravity_Technical_Implementation.md'),
    os.path.join(base, 'AntiGravity_Technical_Implementation.docx'),
    'AntiGravity — Technical Implementation Report'
)

print("Done! Both .docx files created.")
